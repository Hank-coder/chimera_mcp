#!/usr/bin/env python3
"""
文件内容提取器
支持从PDF、Word、Excel等文档中提取文本内容
"""

import asyncio
import tempfile
from pathlib import Path
from typing import Optional, Dict, Any, Tuple
from io import BytesIO

import aiohttp
import aiofiles
from loguru import logger

# 文档处理库
try:
    import PyPDF2
    from docx import Document
    import openpyxl
    DEPENDENCIES_AVAILABLE = True
except ImportError:
    DEPENDENCIES_AVAILABLE = False
    logger.warning("文档处理依赖未安装，文件内容提取功能将被禁用")


class FileContentExtractor:
    """文件内容提取器"""
    
    def __init__(self):
        self.supported_types = {
            'pdf': self._extract_pdf_content,
            'docx': self._extract_docx_content,
            'xlsx': self._extract_xlsx_content,
            'doc': self._extract_doc_content,  # 需要额外处理
            'xls': self._extract_xls_content   # 需要额外处理
        }
        
        # 文件大小限制 (20MB)
        self.max_file_size = 20 * 1024 * 1024
        
        # 内容长度限制 (避免prompt过长)
        self.max_content_length = 8000  # 约8K字符，安全范围
        self.preview_length = 1000      # 预览长度
        
    async def extract_file_content(self, file_url: str, file_type: str, caption: str = "") -> Tuple[str, Dict[str, Any]]:
        """
        从文件URL提取内容
        
        Args:
            file_url: 文件URL
            file_type: 文件类型 (pdf, docx, xlsx等)
            caption: 文件说明
            
        Returns:
            (提取的文本内容, 文件元数据)
        """
        if not DEPENDENCIES_AVAILABLE:
            logger.warning(f"文档处理依赖未安装，跳过文件: {caption}")
            return f"[{file_type.upper()}文件: {caption}] (内容提取不可用)", {}
        
        if file_type.lower() not in self.supported_types:
            logger.info(f"不支持的文件类型: {file_type}")
            return f"[{file_type.upper()}文件: {caption}] (不支持的文件类型)", {}
        
        try:
            # 下载文件
            file_content, file_info = await self._download_file(file_url)
            
            if file_info.get('size', 0) > self.max_file_size:
                logger.warning(f"文件太大，跳过: {caption} ({file_info.get('size', 0)} bytes)")
                return f"[{file_type.upper()}文件: {caption}] (文件过大)", file_info
            
            # 提取内容
            extractor = self.supported_types[file_type.lower()]
            content = await extractor(file_content)
            
            # 处理内容长度，避免prompt过长
            if content:
                original_length = len(content)
                processed_content = self._process_content_length(content, caption)
                
                header = f"[{file_type.upper()}文件: {caption}]\n"
                header += f"文件大小: {file_info.get('size', 0)} bytes\n"
                header += f"内容长度: {original_length} 字符\n"
                
                if len(processed_content) < original_length:
                    header += f"显示: 前 {len(processed_content)} 字符 (已截断)\n"
                
                header += "--- 文件内容 ---\n"
                full_content = header + processed_content + "\n--- 文件内容结束 ---"
            else:
                header = f"[{file_type.upper()}文件: {caption}]\n"
                full_content = header + "(无法提取文本内容)"
            
            file_info['extraction_success'] = bool(content)
            file_info['content_length'] = len(content) if content else 0
            file_info['processed_length'] = len(processed_content) if content else 0
            
            return full_content, file_info
            
        except Exception as e:
            logger.error(f"文件内容提取失败 {caption}: {e}")
            return f"[{file_type.upper()}文件: {caption}] (提取失败: {str(e)})", {}
    
    def _process_content_length(self, content: str, caption: str) -> str:
        """
        智能处理内容长度，避免prompt过长
        
        Args:
            content: 原始内容
            caption: 文件名称
            
        Returns:
            处理后的内容
        """
        if len(content) <= self.max_content_length:
            return content
        
        # 策略1: 智能截断 - 保留开头和重要部分
        lines = content.split('\n')
        processed_lines = []
        current_length = 0
        
        # 优先保留前面的内容（通常包含标题、摘要等重要信息）
        for line in lines:
            line_length = len(line) + 1  # +1 for newline
            if current_length + line_length > self.max_content_length * 0.8:  # 预留20%空间
                break
            processed_lines.append(line)
            current_length += line_length
        
        # 如果还有空间，尝试添加文档结尾的重要信息
        remaining_space = self.max_content_length - current_length
        if remaining_space > 200 and len(lines) > len(processed_lines) + 10:
            # 添加结尾部分
            processed_lines.append("\n... [中间内容已省略] ...")
            
            # 从结尾开始添加行
            end_lines = []
            end_length = 0
            for line in reversed(lines[len(processed_lines):]):
                line_length = len(line) + 1
                if end_length + line_length > remaining_space - 100:  # 预留提示文字空间
                    break
                end_lines.insert(0, line)
                end_length += line_length
            
            if end_lines:
                processed_lines.extend(end_lines)
        
        result = '\n'.join(processed_lines)
        
        # 添加截断提示
        if len(result) < len(content):
            result += f"\n\n[📄 内容已截断: 显示 {len(result)}/{len(content)} 字符，如需完整内容请直接访问文件]"
        
        return result
    
    async def _download_file(self, url: str) -> Tuple[bytes, Dict[str, Any]]:
        """下载文件并返回内容和元数据"""
        async with aiohttp.ClientSession() as session:
            async with session.get(url) as response:
                response.raise_for_status()
                
                file_info = {
                    'size': int(response.headers.get('content-length', 0)),
                    'content_type': response.headers.get('content-type', ''),
                    'url': url
                }
                
                content = await response.read()
                return content, file_info
    
    async def _extract_pdf_content(self, file_content: bytes) -> Optional[str]:
        """提取PDF文本内容"""
        try:
            # 在线程中运行PDF处理（避免阻塞）
            def extract_pdf():
                pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
                text_content = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_content.append(f"--- 第{page_num + 1}页 ---\n{text}")
                    except Exception as e:
                        logger.warning(f"PDF第{page_num + 1}页提取失败: {e}")
                        continue
                
                return "\n\n".join(text_content)
            
            content = await asyncio.get_event_loop().run_in_executor(None, extract_pdf)
            return content if content.strip() else None
            
        except Exception as e:
            logger.error(f"PDF内容提取失败: {e}")
            return None
    
    async def _extract_docx_content(self, file_content: bytes) -> Optional[str]:
        """提取Word文档内容"""
        try:
            def extract_docx():
                doc = Document(BytesIO(file_content))
                text_content = []
                
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        text_content.append(text)
                
                # 提取表格内容
                for table in doc.tables:
                    table_text = []
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells]
                        table_text.append(" | ".join(row_text))
                    if table_text:
                        text_content.append("--- 表格 ---\n" + "\n".join(table_text))
                
                return "\n\n".join(text_content)
            
            content = await asyncio.get_event_loop().run_in_executor(None, extract_docx)
            return content if content.strip() else None
            
        except Exception as e:
            logger.error(f"Word文档内容提取失败: {e}")
            return None
    
    async def _extract_xlsx_content(self, file_content: bytes) -> Optional[str]:
        """提取Excel文档内容"""
        try:
            def extract_xlsx():
                workbook = openpyxl.load_workbook(BytesIO(file_content), read_only=True)
                text_content = []
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    sheet_text = [f"--- 工作表: {sheet_name} ---"]
                    
                    for row in sheet.iter_rows(values_only=True):
                        # 过滤空值并转换为字符串
                        row_values = [str(cell) if cell is not None else "" for cell in row]
                        if any(val.strip() for val in row_values):  # 如果行不是全空
                            sheet_text.append(" | ".join(row_values))
                    
                    if len(sheet_text) > 1:  # 有内容
                        text_content.append("\n".join(sheet_text))
                
                workbook.close()
                return "\n\n".join(text_content)
            
            content = await asyncio.get_event_loop().run_in_executor(None, extract_xlsx)
            return content if content.strip() else None
            
        except Exception as e:
            logger.error(f"Excel文档内容提取失败: {e}")
            return None
    
    async def _extract_doc_content(self, file_content: bytes) -> Optional[str]:
        """提取旧版Word文档内容（.doc格式）"""
        # 旧版Word格式需要额外的库（如python-docx2txt或antiword）
        # 这里先返回占位符
        logger.warning("不支持旧版Word格式(.doc)，请转换为.docx格式")
        return None
    
    async def _extract_xls_content(self, file_content: bytes) -> Optional[str]:
        """提取旧版Excel文档内容（.xls格式）"""
        # 旧版Excel格式需要xlrd库
        logger.warning("不支持旧版Excel格式(.xls)，请转换为.xlsx格式")
        return None


# 全局实例
file_extractor = FileContentExtractor()